"""Minimal Markdown → .docx converter tuned for AeroVertex docs.

Handles: ATX headings, bullets (-), numbered lists, tables (with pipes),
blockquotes (>), fenced code blocks (``` ... ```), horizontal rules (---),
and inline bold (**) / italic (*) / inline-code (`).

Two source files are converted: PROJECT_OVERVIEW.md and PRESENTATION.md.
Output goes next to them in the same docs/ folder.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


HERE = Path(__file__).parent
SOURCES = ["PROJECT_OVERVIEW.md", "PRESENTATION.md"]

ACCENT = RGBColor(0xE6, 0x7E, 0x22)
MUTED = RGBColor(0x66, 0x66, 0x60)
TEXT = RGBColor(0x1A, 0x1A, 0x1A)


# ── Inline span parsing ────────────────────────────────────────────────────
INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*)"           # bold
    r"|(\*[^*]+\*)"              # italic
    r"|(`[^`]+`)"                # inline code
)


def add_inline_runs(paragraph, text: str) -> None:
    """Walk a string and append runs honouring bold/italic/code spans."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ── Block-level helpers ────────────────────────────────────────────────────
def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C8C8C0")
    pbdr.append(bottom)
    p_pr.append(pbdr)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level > 1 else 18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(22)
        run.font.color.rgb = TEXT
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = ACCENT
    elif level == 3:
        run.font.size = Pt(13)
        run.font.color.rgb = TEXT
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = MUTED


def add_paragraph(doc: Document, text: str, *, italic=False, muted=False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if italic or muted:
        # Apply across runs by toggling defaults on a base run, then inline parser
        base = p.add_run("")
        if italic:
            base.italic = True
        if muted:
            base.font.color.rgb = MUTED
    add_inline_runs(p, text)
    if muted:
        for r in p.runs:
            r.font.color.rgb = MUTED


def add_quote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Left border bar
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "E67E22")
    pbdr.append(left)
    p_pr.append(pbdr)
    add_inline_runs(p, text)
    for r in p.runs:
        r.italic = True


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.space_after = Pt(2)
    add_inline_runs(p, text)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    add_inline_runs(p, text)


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    # Light grey shading
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F0F0EB")
    p_pr.append(shd)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.autofit = True

    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        cell = hdr[i]
        set_cell_shading(cell, "F0F0EB")
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h.strip())
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = TEXT

    for r_idx, row in enumerate(rows):
        for c_idx in range(len(header)):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            val = row[c_idx] if c_idx < len(row) else ""
            add_inline_runs(p, val.strip())
            for run in p.runs:
                run.font.size = Pt(10)

    # spacing after the table
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ── Parser ────────────────────────────────────────────────────────────────
def is_table_row(line: str) -> bool:
    return "|" in line and line.strip().startswith("|") and line.strip().endswith("|")


def parse_md(source: str) -> list[tuple[str, object]]:
    """Return a list of (kind, payload) tokens."""
    tokens: list[tuple[str, object]] = []
    lines = source.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Fenced code block
        if stripped.startswith("```"):
            i += 1
            buf = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            tokens.append(("code", buf))
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            tokens.append(("hr", None))
            i += 1
            continue

        # Heading
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            tokens.append(("heading", (len(m.group(1)), m.group(2).strip())))
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip().lstrip(">").strip())
                i += 1
            tokens.append(("quote", " ".join(quote_lines)))
            continue

        # Table (header | --- | row row row)
        if is_table_row(line):
            block = []
            while i < len(lines) and is_table_row(lines[i]):
                block.append(lines[i])
                i += 1
            if len(block) >= 2 and re.search(r"---", block[1]):
                header = [c.strip() for c in block[0].strip().strip("|").split("|")]
                rows = []
                for row_line in block[2:]:
                    cells = [c.strip() for c in row_line.strip().strip("|").split("|")]
                    rows.append(cells)
                tokens.append(("table", (header, rows)))
            else:
                # Not a real table — emit as paragraphs
                for b in block:
                    tokens.append(("para", b))
            continue

        # Bullet
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            indent = len(m.group(1)) // 2
            tokens.append(("bullet", (indent, m.group(2).strip())))
            i += 1
            continue

        # Numbered list
        m = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if m:
            tokens.append(("numbered", m.group(1).strip()))
            i += 1
            continue

        # Blank line
        if stripped == "":
            tokens.append(("blank", None))
            i += 1
            continue

        # Paragraph
        tokens.append(("para", stripped))
        i += 1

    return tokens


# ── Document assembly ────────────────────────────────────────────────────
def build_docx(md_path: Path, docx_path: Path) -> None:
    source = md_path.read_text(encoding="utf-8")
    tokens = parse_md(source)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = TEXT

    for kind, payload in tokens:
        if kind == "heading":
            level, text = payload
            add_heading(doc, text, level)
        elif kind == "para":
            add_paragraph(doc, payload)
        elif kind == "quote":
            add_quote(doc, payload)
        elif kind == "bullet":
            level, text = payload
            add_bullet(doc, text, level)
        elif kind == "numbered":
            add_numbered(doc, payload)
        elif kind == "code":
            add_code_block(doc, payload)
        elif kind == "table":
            header, rows = payload
            add_table(doc, header, rows)
        elif kind == "hr":
            add_horizontal_rule(doc)
        elif kind == "blank":
            continue

    doc.save(docx_path)
    print(f"  wrote {docx_path.name} ({docx_path.stat().st_size:,} bytes)")


if __name__ == "__main__":
    for md_name in SOURCES:
        md_path = HERE / md_name
        if not md_path.exists():
            print(f"  skip {md_name} (not found)")
            continue
        docx_path = HERE / md_name.replace(".md", ".docx")
        print(f"converting {md_name} -> {docx_path.name}")
        build_docx(md_path, docx_path)
    print("done.")
